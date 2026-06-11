from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Literal

from bs4 import BeautifulSoup

from dean_os.research_corpus import ResearchCorpus
from dean_os.schemas import ResearchDocument

ResearchSourceType = Literal["news", "article", "book", "report", "filing", "transcript"]

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".json", ".pdf", ".docx"}


class MaterialLoadError(RuntimeError):
    pass


def load_research_document(
    path: str | Path,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
) -> ResearchDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise MaterialLoadError(f"Unsupported research material extension: {suffix}")
    if not file_path.exists():
        raise MaterialLoadError(f"Research material does not exist: {file_path}")

    if suffix in {".txt", ".md", ".markdown"}:
        raw_text = _read_text_file(file_path)
        title = _title_from_markdown(raw_text) or file_path.stem
        text = _clean_text(raw_text)
    elif suffix in {".html", ".htm"}:
        raw_html = _read_text_file(file_path)
        title, text = _extract_html(raw_html, fallback_title=file_path.stem)
    elif suffix == ".json":
        title, text, metadata = _extract_json(file_path)
        return ResearchDocument(
            title=title,
            source_type=source_type or _infer_source_type(file_path),
            text=text,
            uri=str(file_path),
            tickers=tickers or metadata.get("tickers", []),
            sectors=sectors or metadata.get("sectors", []),
            tags=tags or metadata.get("tags", []),
            published_at=metadata.get("published_at"),
            metadata={"path": str(file_path), **metadata},
        )
    elif suffix == ".pdf":
        title = file_path.stem
        text = _extract_pdf(file_path)
    elif suffix == ".docx":
        title = file_path.stem
        text = _extract_docx(file_path)
    else:
        raise MaterialLoadError(f"Unsupported research material extension: {suffix}")

    if not text.strip():
        raise MaterialLoadError(f"No extractable text found in: {file_path}")

    return ResearchDocument(
        title=title,
        source_type=source_type or _infer_source_type(file_path),
        text=text,
        uri=str(file_path),
        tickers=tickers or [],
        sectors=sectors or [],
        tags=tags or [],
        metadata={"path": str(file_path), "extension": suffix},
    )


def load_research_directory(
    path: str | Path,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
    recursive: bool = True,
    ignore_errors: bool = True,
) -> tuple[list[ResearchDocument], list[str]]:
    root = Path(path)
    if root.is_file():
        return [load_research_document(root, source_type, tickers, sectors, tags)], []
    if not root.exists():
        raise MaterialLoadError(f"Research material path does not exist: {root}")

    pattern = "**/*" if recursive else "*"
    documents: list[ResearchDocument] = []
    errors: list[str] = []
    for file_path in sorted(root.glob(pattern)):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            documents.append(load_research_document(file_path, source_type, tickers, sectors, tags))
        except MaterialLoadError as exc:
            if not ignore_errors:
                raise
            errors.append(str(exc))
    return documents, errors


def ingest_research_path(
    path: str | Path,
    corpus: ResearchCorpus,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
    chunk_size: int = 1200,
    recursive: bool = True,
) -> dict:
    documents, errors = load_research_directory(
        path=path,
        source_type=source_type,
        tickers=tickers,
        sectors=sectors,
        tags=tags,
        recursive=recursive,
        ignore_errors=True,
    )
    chunk_count = 0
    for document in documents:
        chunk_count += len(corpus.add_document(document, chunk_size=chunk_size))
    return {
        "document_count": len(documents),
        "chunk_count": chunk_count,
        "errors": errors,
        "titles": [document.title for document in documents],
    }


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise MaterialLoadError(f"Could not decode text file: {path}")


def _title_from_markdown(text: str) -> str | None:
    for line in text.splitlines():
        match = re.match(r"^\s*#\s+(.+?)\s*$", line)
        if match:
            return match.group(1).strip()
    return None


def _extract_html(raw_html: str, fallback_title: str) -> tuple[str, str]:
    soup = BeautifulSoup(raw_html, "lxml")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    title = soup.title.get_text(" ", strip=True) if soup.title else fallback_title
    text = soup.get_text(" ", strip=True)
    return title or fallback_title, _clean_text(text)


def _extract_json(path: Path) -> tuple[str, str, dict]:
    try:
        payload = json.loads(_read_text_file(path))
    except json.JSONDecodeError as exc:
        raise MaterialLoadError(f"Invalid JSON research material: {path}") from exc
    if isinstance(payload, list):
        text = "\n".join(_text_from_json_item(item) for item in payload)
        return path.stem, _clean_text(text), {"item_count": len(payload)}
    if not isinstance(payload, dict):
        return path.stem, _clean_text(str(payload)), {}
    title = str(payload.get("title") or payload.get("headline") or path.stem)
    text = _text_from_json_item(payload)
    metadata = {key: value for key, value in payload.items() if key not in {"text", "content", "body", "summary"}}
    return title, _clean_text(text), metadata


def _text_from_json_item(item) -> str:
    if not isinstance(item, dict):
        return str(item)
    return " ".join(
        str(item.get(key, ""))
        for key in ("title", "headline", "summary", "description", "content", "body", "text")
        if item.get(key)
    )


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise MaterialLoadError("PDF ingestion requires optional package pypdf or PyPDF2") from exc
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n".join(pages))


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise MaterialLoadError("DOCX ingestion requires optional package python-docx") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return _clean_text("\n".join(paragraphs))


def _infer_source_type(path: Path) -> ResearchSourceType:
    lower_name = path.name.lower()
    if "10-k" in lower_name or "10-q" in lower_name or "filing" in lower_name:
        return "filing"
    if "transcript" in lower_name or "earnings-call" in lower_name:
        return "transcript"
    if "book" in lower_name or "chapter" in lower_name:
        return "book"
    if "news" in lower_name:
        return "news"
    if "report" in lower_name or path.suffix.lower() in {".pdf", ".docx"}:
        return "report"
    return "article"


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
