from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Literal

from bs4 import BeautifulSoup

from dean_os.draft.dean_os_agent_system_v7.dean_os.research_corpus import ResearchCorpus
from dean_os.schemas import ResearchDocument

ResearchSourceType = Literal["news", "article", "book", "report", "filing", "transcript"]

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".json", ".pdf", ".docx"}
SENTIMENT_EXCLUDED_QUARANTINE_FLAGS = frozenset(
    {"legal_disclaimer", "third_party_rating", "advertising_navigation_author_bio"}
)

QUARANTINE_PATTERNS: tuple[tuple[str, tuple[re.Pattern[str], ...]], ...] = (
    (
        "legal_disclaimer",
        (
            re.compile(r"\bforward-looking statements?\b", re.IGNORECASE),
            re.compile(r"\bsafe harbor\b", re.IGNORECASE),
            re.compile(r"\bactual results may differ\b", re.IGNORECASE),
            re.compile(r"\bundertakes? no obligation to update\b", re.IGNORECASE),
            re.compile(r"\bno obligation to update\b", re.IGNORECASE),
            re.compile(r"\bnot (?:financial|investment) advice\b", re.IGNORECASE),
            re.compile(r"\bnot a recommendation\b", re.IGNORECASE),
            re.compile(r"\bfor informational purposes only\b", re.IGNORECASE),
            re.compile(r"\brisks and uncertainties\b", re.IGNORECASE),
            re.compile(r"\brisk factors\b.*\bsec filings\b", re.IGNORECASE),
            re.compile(r"\blegal notice\b", re.IGNORECASE),
        ),
    ),
    (
        "third_party_rating",
        (
            re.compile(r"\bthird[- ]party rating\b", re.IGNORECASE),
            re.compile(r"\banalyst rating\b", re.IGNORECASE),
            re.compile(r"\bbroker rating\b", re.IGNORECASE),
            re.compile(r"\bconsensus rating\b", re.IGNORECASE),
            re.compile(r"\bprice target\b", re.IGNORECASE),
            re.compile(r"\bmaintains?\s+(?:buy|sell|hold|outperform|underperform)\s+rating\b", re.IGNORECASE),
            re.compile(r"\b(?:upgrades?|downgrades?)\s+(?:to|from)\s+(?:buy|sell|hold|outperform|underperform)\b", re.IGNORECASE),
            re.compile(r"\bzacks rank\b", re.IGNORECASE),
            re.compile(r"\bmorningstar rating\b", re.IGNORECASE),
            re.compile(r"\btipranks\b", re.IGNORECASE),
        ),
    ),
    (
        "advertising_navigation_author_bio",
        (
            re.compile(r"\badvertisement\b", re.IGNORECASE),
            re.compile(r"\bsponsored content\b", re.IGNORECASE),
            re.compile(r"\bsubscribe now\b", re.IGNORECASE),
            re.compile(r"\bsign up\b.*\bnewsletter\b", re.IGNORECASE),
            re.compile(r"\bclick here to\b", re.IGNORECASE),
            re.compile(r"\bcookie policy\b", re.IGNORECASE),
            re.compile(r"\bprivacy policy\b", re.IGNORECASE),
            re.compile(r"\bterms of use\b", re.IGNORECASE),
            re.compile(r"\ball rights reserved\b", re.IGNORECASE),
            re.compile(r"\babout the author\b", re.IGNORECASE),
            re.compile(r"\bauthor bio\b", re.IGNORECASE),
            re.compile(r"\bfollow us\b", re.IGNORECASE),
            re.compile(r"\brelated articles\b", re.IGNORECASE),
        ),
    ),
)

_PARAGRAPH_RE = re.compile(r"\S[\s\S]*?(?=\n\s*\n|\Z)")
_SENTENCE_RE = re.compile(r"[^.!?\n]+(?:[.!?]+|$)")


class MaterialLoadError(RuntimeError):
    pass


def load_research_document(
    path: str | Path,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
) -> ResearchDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise MaterialLoadError(f"Unsupported research material extension: {suffix}")
    if not file_path.exists():
        raise MaterialLoadError(f"Research material does not exist: {file_path}")

    if suffix in {".txt", ".md", ".markdown"}:
        raw_text = _read_text_file(file_path)
        title = _title_from_markdown(raw_text) or file_path.stem
        text = _clean_text(raw_text)
    elif suffix in {".html", ".htm"}:
        raw_html = _read_text_file(file_path)
        title, text = _extract_html(raw_html, fallback_title=file_path.stem)
    elif suffix == ".json":
        title, text, metadata = _extract_json(file_path)
        return annotate_quarantine(
            ResearchDocument(
                title=title,
                source_type=source_type or _infer_source_type(file_path),
                text=text,
                uri=str(file_path),
                tickers=tickers or metadata.get("tickers", []),
                sectors=sectors or metadata.get("sectors", []),
                tags=tags or metadata.get("tags", []),
                published_at=metadata.get("published_at"),
                metadata={"path": str(file_path), **metadata},
            )
        )
    elif suffix == ".pdf":
        title = file_path.stem
        text = _extract_pdf(file_path)
    elif suffix == ".docx":
        title = file_path.stem
        text = _extract_docx(file_path)
    else:
        raise MaterialLoadError(f"Unsupported research material extension: {suffix}")

    if not text.strip():
        raise MaterialLoadError(f"No extractable text found in: {file_path}")

    return annotate_quarantine(
        ResearchDocument(
            title=title,
            source_type=source_type or _infer_source_type(file_path),
            text=text,
            uri=str(file_path),
            tickers=tickers or [],
            sectors=sectors or [],
            tags=tags or [],
            metadata={"path": str(file_path), "extension": suffix},
        )
    )


def load_research_directory(
    path: str | Path,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
    recursive: bool = True,
    ignore_errors: bool = True,
) -> tuple[list[ResearchDocument], list[str]]:
    root = Path(path)
    if root.is_file():
        return [load_research_document(root, source_type, tickers, sectors, tags)], []
    if not root.exists():
        raise MaterialLoadError(f"Research material path does not exist: {root}")

    pattern = "**/*" if recursive else "*"
    documents: list[ResearchDocument] = []
    errors: list[str] = []
    for file_path in sorted(root.glob(pattern)):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            documents.append(load_research_document(file_path, source_type, tickers, sectors, tags))
        except MaterialLoadError as exc:
            if not ignore_errors:
                raise
            errors.append(str(exc))
    return documents, errors


def ingest_research_path(
    path: str | Path,
    corpus: ResearchCorpus,
    source_type: ResearchSourceType | None = None,
    tickers: list[str] | None = None,
    sectors: list[str] | None = None,
    tags: list[str] | None = None,
    chunk_size: int = 1200,
    recursive: bool = True,
) -> dict:
    documents, errors = load_research_directory(
        path=path,
        source_type=source_type,
        tickers=tickers,
        sectors=sectors,
        tags=tags,
        recursive=recursive,
        ignore_errors=True,
    )
    chunk_count = 0
    for document in documents:
        chunk_count += len(corpus.add_document(document, chunk_size=chunk_size))
    return {
        "document_count": len(documents),
        "chunk_count": chunk_count,
        "errors": errors,
        "titles": [document.title for document in documents],
    }


def _read_text_file(path: Path) -> str:
    from dean_os.draft.dean_os_agent_system_v7.dean_os.dean_paths import DeanPaths

    return DeanPaths.load_text_file(path)


def annotate_quarantine(document: ResearchDocument) -> ResearchDocument:
    """Attach document-level quarantine metadata without changing source text."""
    quarantine_blocks = detect_quarantine_blocks(document.text)
    quarantine_flags = sorted(
        {
            flag
            for block in quarantine_blocks
            for flag in block.get("quarantine_flags", [])
        }
    )
    quality_precheck = "quarantine_detected" if quarantine_flags else "passed"
    metadata = {
        **document.metadata,
        "quarantine_blocks": quarantine_blocks,
        "quarantine_block_count": len(quarantine_blocks),
        "quarantine_flags": quarantine_flags,
        "quality_precheck": quality_precheck,
    }
    return document.model_copy(
        update={
            "metadata": metadata,
            "quarantine_flags": quarantine_flags,
            "quality_precheck": quality_precheck,
        }
    )


def detect_quarantine_blocks(text: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for block_index, partition in enumerate(partition_text_by_quarantine(text)):
        flags = partition["quarantine_flags"]
        if not flags:
            continue
        block_text = partition["text"]
        blocks.append(
            {
                "block_index": block_index,
                "start": partition["start"],
                "end": partition["end"],
                "quarantine_flags": flags,
                "text_preview": block_text[:240],
            }
        )
    return blocks


def quarantine_flags_for_text(text: str) -> list[str]:
    flags: list[str] = []
    for flag, patterns in QUARANTINE_PATTERNS:
        if any(pattern.search(text) for pattern in patterns):
            flags.append(flag)
    return flags


def partition_text_by_quarantine(text: str) -> list[dict[str, Any]]:
    partitions: list[dict[str, Any]] = []
    for unit in _analysis_units(text):
        unit_text = _clean_text(unit["text"])
        if not unit_text:
            continue
        flags = quarantine_flags_for_text(unit_text)
        if partitions and partitions[-1]["quarantine_flags"] == flags:
            partitions[-1]["text"] = f"{partitions[-1]['text']} {unit_text}".strip()
            partitions[-1]["end"] = unit["end"]
            continue
        partitions.append(
            {
                "text": unit_text,
                "start": unit["start"],
                "end": unit["end"],
                "quarantine_flags": flags,
                "quality_precheck": "quarantined" if flags else "passed",
            }
        )
    return partitions


def filter_quarantined_text(
    text: str,
    excluded_flags: set[str] | frozenset[str] = SENTIMENT_EXCLUDED_QUARANTINE_FLAGS,
) -> tuple[str, list[dict[str, Any]]]:
    kept: list[str] = []
    removed: list[dict[str, Any]] = []
    for partition in partition_text_by_quarantine(text):
        flags = set(partition["quarantine_flags"])
        if flags.intersection(excluded_flags):
            removed.append(partition)
            continue
        kept.append(partition["text"])
    return _clean_text(" ".join(kept)), removed


def sentiment_safe_text(document: ResearchDocument) -> str:
    safe_text, _removed = filter_quarantined_text(document.text)
    return safe_text


def _analysis_units(text: str) -> list[dict[str, Any]]:
    if not text:
        return []
    paragraphs = [
        {"text": match.group(0), "start": match.start(), "end": match.end()}
        for match in _PARAGRAPH_RE.finditer(text)
        if match.group(0).strip()
    ]
    if len(paragraphs) > 1:
        return paragraphs

    units: list[dict[str, Any]] = []
    for match in _SENTENCE_RE.finditer(text):
        sentence = match.group(0)
        if sentence.strip():
            units.append({"text": sentence, "start": match.start(), "end": match.end()})
    if units:
        return units
    return [{"text": text, "start": 0, "end": len(text)}]


def _title_from_markdown(text: str) -> str | None:
    for line in text.splitlines():
        match = re.match(r"^\s*#\s+(.+?)\s*$", line)
        if match:
            return match.group(1).strip()
    return None


def _extract_html(raw_html: str, fallback_title: str) -> tuple[str, str]:
    soup = BeautifulSoup(raw_html, "lxml")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    title = soup.title.get_text(" ", strip=True) if soup.title else fallback_title
    text = soup.get_text(" ", strip=True)
    return title or fallback_title, _clean_text(text)


def _extract_json(path: Path) -> tuple[str, str, dict]:
    try:
        payload = json.loads(_read_text_file(path))
    except json.JSONDecodeError as exc:
        raise MaterialLoadError(f"Invalid JSON research material: {path}") from exc
    if isinstance(payload, list):
        text = "\n".join(_text_from_json_item(item) for item in payload)
        return path.stem, _clean_text(text), {"item_count": len(payload)}
    if not isinstance(payload, dict):
        return path.stem, _clean_text(str(payload)), {}
    title = str(payload.get("title") or payload.get("headline") or path.stem)
    text = _text_from_json_item(payload)
    metadata = {key: value for key, value in payload.items() if key not in {"text", "content", "body", "summary"}}
    return title, _clean_text(text), metadata


def _text_from_json_item(item) -> str:
    if not isinstance(item, dict):
        return str(item)
    return " ".join(
        str(item.get(key, ""))
        for key in ("title", "headline", "summary", "description", "content", "body", "text")
        if item.get(key)
    )


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise MaterialLoadError("PDF ingestion requires optional package pypdf or PyPDF2") from exc
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n".join(pages))


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise MaterialLoadError("DOCX ingestion requires optional package python-docx") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return _clean_text("\n".join(paragraphs))


def _infer_source_type(path: Path) -> ResearchSourceType:
    lower_name = path.name.lower()
    if "10-k" in lower_name or "10-q" in lower_name or "filing" in lower_name:
        return "filing"
    if "transcript" in lower_name or "earnings-call" in lower_name:
        return "transcript"
    if "book" in lower_name or "chapter" in lower_name:
        return "book"
    if "news" in lower_name:
        return "news"
    if "report" in lower_name or path.suffix.lower() in {".pdf", ".docx"}:
        return "report"
    return "article"


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
